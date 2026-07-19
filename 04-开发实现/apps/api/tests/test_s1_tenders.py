from __future__ import annotations

from datetime import date
from html import unescape
from pathlib import Path
import re
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.core.api_response import ApiError
from app.core.audit import AuditRepository
from app.core.config import Settings
from app.main import create_app
from app.modules.accounts.repository import InMemoryAuthRepository
from app.modules.platform.append_only_storage import LocalAppendOnlyFileStorage
from app.modules.steven.tender_application import StevenTenderApplicationService
from app.modules.steven.tender_repository import DEMO_TEMPLATES, InMemoryTenderRepository
from app.modules.steven.tender_schemas import TenderCreateRequest, TenderUpdateRequest
from app.modules.steven.tender_uow import InMemoryTenderUnitOfWork
from app.modules.steven.tender_word import TenderWordRenderer


def payload(**overrides):
    values = {
        "template_id": "template-s1-demo-001",
        "title": "脱敏采购服务邀请文书",
        "document_number": "DEMO-S1-20260717-001",
        "subject": "脱敏设施保养服务",
        "generated_date": "2026-07-17",
        "deadline_date": "2026-07-20",
        "budget_min": 1000,
        "budget_max": 3000,
        "currency": "HKD",
        "location": "演示地点 A（脱敏）",
        "controlled_clauses": "仅接受脱敏演示资料；最终内容必须由人工复核。",
        "supplier_names": ["演示供应商甲", "演示供应商乙", "演示供应商丙"],
        "is_demo": True,
    }
    values.update(overrides)
    return values


def make_application(tmp_path):
    repository = InMemoryTenderRepository(seed_demo=True)
    audit = AuditRepository()
    renderer = TenderWordRenderer()
    application = StevenTenderApplicationService(
        InMemoryTenderUnitOfWork(repository, audit),
        renderer,
        LocalAppendOnlyFileStorage(tmp_path, "tenders", "docx", renderer.verify),
    )
    return application, repository, audit


def create_tender(application, **overrides):
    return application.create_tender(TenderCreateRequest(**payload(**overrides)), "steven.test")


def approve_tender(application, tender_id: str):
    application.preview(tender_id, "steven.test")
    application.submit(tender_id, "steven.test")
    return application.approve(tender_id, "approver.test", "已由独立审批人完成人工复核")


def docx_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    parts.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return "\n".join(parts)


def test_three_distinct_demo_templates_are_seeded_and_renderable(tmp_path):
    application, _, _ = make_application(tmp_path)
    templates = application.list_templates()
    assert len(templates) >= 3
    assert {item.document_type for item in templates}.issuperset({item.document_type for item in DEMO_TEMPLATES})
    assert len({tuple(item.variables) for item in templates}) >= 3

    for index, template in enumerate(templates[:3], start=1):
        tender = create_tender(
            application,
            template_id=template.id,
            document_number=f"DEMO-S1-TEMPLATE-{index}",
        )
        preview = application.preview(tender.id, "steven.test")
        assert preview.valid is True
        assert "{{" not in (preview.tender.rendered_body or "")


def test_template_live_preview_accepts_unsaved_draft_fields():
    client = TestClient(
        create_app(
            Settings(app_env="test", auth_mode="mock", demo_seed_enabled=True, mock_identity="steven"),
            InMemoryAuthRepository(),
        )
    )
    response = client.post(
        "/api/v1/steven/tender-templates/template-s1-demo-001/preview",
        json={
            "title": "即时预览标题（脱敏）",
            "document_number": "DEMO-S1-LIVE-PREVIEW",
            "subject": "即时预览事项（脱敏）",
            "generated_date": "2026-07-18",
            "deadline_date": "2026-07-21",
            "budget_min": 1200,
            "budget_max": 3600,
            "currency": "HKD",
            "location": "演示地点 B（脱敏）",
            "controlled_clauses": "仅用于本机脱敏预览。",
            "supplier_names": ["演示供应商甲", "演示供应商乙"],
        },
    )

    assert response.status_code == 200
    preview = response.json()["data"]
    assert preview["template_id"] == "template-s1-demo-001"
    assert "即时预览标题（脱敏）" in preview["html"]
    assert "演示供应商甲" in preview["html"]


def test_create_preview_and_draft_docx_use_same_structure(tmp_path):
    application, _, _ = make_application(tmp_path)
    tender = create_tender(application)
    preview = application.preview(tender.id, "steven.test")
    assert preview.valid is True
    assert preview.tender.status == "draft"
    assert preview.tender.unresolved_variables == []
    assert "demo-contact@example.invalid" in (preview.tender.rendered_body or "")

    filename, content = application.draft_bytes(tender.id, "steven.test")
    path = tmp_path / filename
    path.write_bytes(content)
    document = Document(path)
    full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "草稿／未批准／仅供人工复核" in full_text
    assert "当前使用脱敏演示资料" in "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    TenderWordRenderer.verify(path)


@pytest.mark.parametrize(
    ("overrides", "expected_text"),
    [
        ({"deadline_date": "2026-07-19"}, "截止日期"),
        ({"budget_min": -1}, "greater_than_equal"),
        ({"budget_min": 3001, "budget_max": 3000}, "预算下限"),
    ],
)
def test_api_blocks_invalid_dates_and_budgets(overrides, expected_text):
    client = TestClient(
        create_app(
            Settings(app_env="test", auth_mode="mock", demo_seed_enabled=True, mock_identity="steven"),
            InMemoryAuthRepository(),
        )
    )
    response = client.post("/api/v1/steven/tenders", json=payload(**overrides))
    assert response.status_code == 422
    assert expected_text in str(response.json())


@pytest.mark.parametrize(
    "supplier_names",
    [
        ["Demo Vendor", "demo vendor"],
        ["ＡＢＣ 供应商", "ABC 供应商"],
        ["演示　供应商", " 演示  供应商 "],
    ],
)
def test_nfkc_casefold_and_whitespace_duplicate_suppliers_are_blocked(tmp_path, supplier_names):
    application, _, _ = make_application(tmp_path)
    with pytest.raises(ApiError) as captured:
        create_tender(application, supplier_names=supplier_names)
    assert captured.value.code == "duplicate_supplier_name"


def test_unresolved_variable_persists_draft_error_and_blocks_submit(tmp_path):
    application, repository, audit = make_application(tmp_path)
    repository.templates["template-s1-demo-001"].template_body += "\n待补字段：{{missing_demo_value}}"
    tender = create_tender(application)
    preview = application.preview(tender.id, "steven.test")
    assert preview.valid is False
    assert preview.tender.status == "draft_error"
    assert preview.unresolved_variables == ["{{missing_demo_value}}"]
    assert application.get_tender(tender.id).status == "draft_error"

    with pytest.raises(ApiError) as captured:
        application.submit(tender.id, "steven.test")
    assert captured.value.code == "draft_not_ready"
    assert application.get_tender(tender.id).status == "draft_error"
    assert any(event["action"] == "tender.submit_rejected" for event in audit.list_for_object(tender.id))


def test_submitter_cannot_self_approve_and_independent_approver_can(tmp_path):
    application, _, _ = make_application(tmp_path)
    tender = create_tender(application)
    application.preview(tender.id, "dual-role.test")
    application.submit(tender.id, "dual-role.test")
    with pytest.raises(ApiError) as captured:
        application.approve(tender.id, "dual-role.test", "不应允许自审")
    assert captured.value.status_code == 403
    assert captured.value.code == "self_approval_forbidden"
    approved = application.approve(tender.id, "approver.test", "独立审批人已核对")
    assert approved.status == "approved"


def test_return_requires_opinion_and_returned_tender_cannot_export(tmp_path):
    application, _, _ = make_application(tmp_path)
    tender = create_tender(application)
    application.preview(tender.id, "steven.test")
    application.submit(tender.id, "steven.test")
    with pytest.raises(ApiError) as captured:
        application.return_for_revision(tender.id, "approver.test", " ")
    assert captured.value.code == "return_opinion_required"
    returned = application.return_for_revision(tender.id, "approver.test", "请补充服务范围说明")
    assert returned.status == "returned"
    with pytest.raises(ApiError) as export_error:
        application.export(tender.id, "steven.test")
    assert export_error.value.code == "formal_export_forbidden"


@pytest.mark.parametrize("state", ["draft", "draft_error", "submitted", "returned"])
def test_non_approved_states_cannot_formally_export(tmp_path, state):
    application, repository, _ = make_application(tmp_path)
    tender = create_tender(application, document_number=f"DEMO-S1-{state}")
    if state == "draft_error":
        repository.templates["template-s1-demo-001"].template_body += "\n{{missing_value}}"
        application.preview(tender.id, "steven.test")
    elif state == "submitted":
        application.preview(tender.id, "steven.test")
        application.submit(tender.id, "steven.test")
    elif state == "returned":
        application.preview(tender.id, "steven.test")
        application.submit(tender.id, "steven.test")
        application.return_for_revision(tender.id, "approver.test", "退回测试")
    with pytest.raises(ApiError) as captured:
        application.export(tender.id, "steven.test")
    assert captured.value.code == "formal_export_forbidden"


def test_approved_exports_are_append_only_reopenable_and_audited(tmp_path):
    application, repository, audit = make_application(tmp_path)
    tender = create_tender(application)
    approve_tender(application, tender.id)

    exports = [application.export(tender.id, "steven.test") for _ in range(3)]
    assert [item.version.version_number for item in exports] == [1, 2, 3]
    versions = application.list_versions(tender.id)
    assert [item.version_number for item in versions] == [1, 2, 3]
    assert len({item.storage_key for item in versions}) == 3
    assert len({item.sha256 for item in versions}) == 3

    for version in versions:
        path = application.version_file(tender.id, version.version_number)
        document = Document(path)
        assert document.paragraphs[0].text == "正式批准版本"
        assert version.size_bytes == path.stat().st_size
        assert version.sha256 is not None and len(version.sha256) == 64
        assert repository.files[version.file_id]["request_id"] is None

    actions = [event["action"] for event in audit.list_for_object(tender.id)]
    assert {"tender.create", "tender.preview", "tender.submit", "tender.approve", "tender.export_reserved", "tender.export"} <= set(actions)


def test_approved_batch_export_creates_supplier_isolated_docx_files(tmp_path):
    application, _, audit = make_application(tmp_path)
    tender = create_tender(
        application,
        template_id="template-s1-demo-002",
        document_number="DEMO-S1-BATCH-001",
        supplier_names=["演示供应商甲", "演示供应商乙", "演示供应商丙"],
        controlled_clauses="每份报价必须按对应受邀供应商独立提交；不得转交其他供应商。",
    )
    approved = approve_tender(application, tender.id)
    selected = approved.suppliers[:2]
    result = application.batch_export(tender.id, [item.id for item in selected], "steven.test")

    assert len(result.versions) == 2
    assert {item.export_batch_id for item in result.versions} == {result.batch_id}
    assert {item.supplier_id for item in result.versions} == {item.id for item in selected}
    assert {item.supplier_name_snapshot for item in result.versions} == {item.supplier_name for item in selected}
    assert len({item.storage_key for item in result.versions}) == 2

    for version in result.versions:
        path = application.version_file(tender.id, version.version_number)
        text = docx_text(path)
        own_name = version.supplier_name_snapshot or ""
        other_names = {item.supplier_name for item in selected} - {own_name}
        assert own_name in text
        assert all(other_name not in text for other_name in other_names)
        assert "不得转交其他供应商" in text
        assert "{{" not in text
        TenderWordRenderer.verify(path)

    actions = [event["action"] for event in audit.list_for_object(tender.id)]
    assert "tender.batch_export_reserved" in actions
    assert "tender.batch_export" in actions


def test_batch_export_rejects_invalid_supplier_selection(tmp_path):
    application, _, _ = make_application(tmp_path)
    tender = create_tender(application, document_number="DEMO-S1-BATCH-BLOCK-001")
    with pytest.raises(ApiError, match="仅已批准文书"):
        application.batch_export(tender.id, [item.id for item in tender.suppliers[:2]], "steven.test")

    approved = approve_tender(application, tender.id)
    with pytest.raises(ApiError, match="2 至 20"):
        application.batch_export(tender.id, [approved.suppliers[0].id], "steven.test")
    with pytest.raises(ApiError, match="不可重复"):
        application.batch_export(tender.id, [approved.suppliers[0].id, approved.suppliers[0].id], "steven.test")
    with pytest.raises(ApiError, match="属于当前文书事项"):
        application.batch_export(tender.id, [approved.suppliers[0].id, "outside-supplier"], "steven.test")


def test_manual_rendered_body_revision_persists_and_keeps_unresolved_markers_blocking(tmp_path):
    application, _, _ = make_application(tmp_path)
    tender = create_tender(application, document_number="DEMO-S1-MANUAL-BODY-001")
    preview = application.preview(tender.id, "steven.test")
    original_body = preview.tender.rendered_body or ""

    manually_revised = original_body.replace(
        "仅接受脱敏演示资料；最终内容必须由人工复核。",
        "仅接受脱敏演示资料，最终内容必须由人工复核。",
    )
    updated = application.update_tender(
        tender.id,
        TenderUpdateRequest(rendered_body=manually_revised),
        "steven.test",
    )
    assert updated.rendered_body == manually_revised
    assert updated.status == "draft"
    assert updated.unresolved_variables == []

    unresolved_body = f"{manually_revised}\n待人工补充：[待人工輸入]"
    blocked = application.update_tender(
        tender.id,
        TenderUpdateRequest(rendered_body=unresolved_body),
        "steven.test",
    )
    assert blocked.status == "draft_error"
    assert blocked.unresolved_variables == ["[待人工輸入]"]
    with pytest.raises(ApiError) as captured:
        application.submit(tender.id, "steven.test")
    assert captured.value.code == "draft_not_ready"

    resolved_body = unresolved_body.replace("[待人工輸入]", "脱敏人工补充值")
    resolved = application.update_tender(
        tender.id,
        TenderUpdateRequest(rendered_body=resolved_body),
        "steven.test",
    )
    assert resolved.status == "draft"
    assert resolved.unresolved_variables == []
    assert resolved.rendered_body == resolved_body



def test_update_after_return_resets_approval_state(tmp_path):
    application, _, _ = make_application(tmp_path)
    tender = create_tender(application)
    application.preview(tender.id, "steven.test")
    application.submit(tender.id, "steven.test")
    application.return_for_revision(tender.id, "approver.test", "请修订条款")
    updated = application.update_tender(
        tender.id,
        TenderUpdateRequest(controlled_clauses="已按人工退回意见修订的脱敏条款。"),
        "steven.test",
    )
    assert updated.status == "draft"
    assert updated.submitted_by is None
    assert updated.decided_by is None
    assert updated.decision_opinion is None


def test_api_permission_matrix_and_request_id_audit():
    steven = TestClient(
        create_app(
            Settings(app_env="test", auth_mode="mock", demo_seed_enabled=True, mock_identity="steven"),
            InMemoryAuthRepository(),
        )
    )
    assert steven.get("/api/v1/steven/tender-templates").status_code == 200
    created = steven.post(
        "/api/v1/steven/tenders",
        json=payload(document_number="DEMO-S1-AUDIT-001"),
        headers={"X-Request-Id": "s1-request-001"},
    )
    assert created.status_code == 201
    tender_id = created.json()["data"]["id"]
    audit_response = steven.get(f"/api/v1/steven/tenders/{tender_id}/audit-events")
    assert audit_response.status_code == 200
    assert any(event["action"] == "tender.create" and event["request_id"] == "s1-request-001" for event in audit_response.json()["data"])
    assert steven.post(f"/api/v1/steven/tenders/{tender_id}/approve", json={"opinion": "越权"}).status_code == 403

    approver = TestClient(
        create_app(
            Settings(app_env="test", auth_mode="mock", demo_seed_enabled=True, mock_identity="approver"),
            InMemoryAuthRepository(),
        )
    )
    assert approver.get("/api/v1/steven/tender-templates").status_code == 200
    assert approver.post("/api/v1/steven/tenders", json=payload(document_number="DEMO-S1-DENIED")).status_code == 403

    admin = TestClient(
        create_app(
            Settings(app_env="test", auth_mode="mock", demo_seed_enabled=True, mock_identity="admin"),
            InMemoryAuthRepository(),
        )
    )
    assert admin.get("/api/v1/steven/tenders").status_code == 403


def test_docx_verifier_rejects_active_content_and_external_relationships(tmp_path):
    renderer = TenderWordRenderer()
    application, _, _ = make_application(tmp_path)
    tender = create_tender(application)
    preview = application.preview(tender.id, "steven.test")
    job = application._run(lambda service: service._require_job(tender.id), tender_id=tender.id)
    suppliers = application._run(lambda service: service.repository.suppliers_for(tender.id), tender_id=tender.id)

    macro_path = tmp_path / "macro.docx"
    renderer.render_to(macro_path, job=job, suppliers=suppliers, rendered_body=preview.tender.rendered_body or "", formal=False)
    with ZipFile(macro_path, "a", ZIP_DEFLATED) as package:
        package.writestr("word/vbaProject.bin", b"not-a-real-macro")
    with pytest.raises(ValueError, match="forbidden_docx_active_content"):
        renderer.verify(macro_path)

    external_path = tmp_path / "external.docx"
    renderer.render_to(external_path, job=job, suppliers=suppliers, rendered_body=preview.tender.rendered_body or "", formal=False)
    with ZipFile(external_path, "a", ZIP_DEFLATED) as package:
        package.writestr(
            "customXml/_rels/external.rels",
            '<?xml version="1.0" encoding="UTF-8"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" Type="demo" Target="https://example.invalid" TargetMode="External"/>'
            "</Relationships>",
        )
    with pytest.raises(ValueError, match="forbidden_docx_external_relationship"):
        renderer.verify(external_path)


def test_0008_migration_appends_s1_tables_constraints_and_permissions():
    migration = (
        Path(__file__).resolve().parents[1]
        / "alembic"
        / "versions"
        / "20260717_0008_steven_s1_tenders.py"
    ).read_text(encoding="utf-8")
    assert 'revision = "20260717_0008"' in migration
    assert 'down_revision = "20260717_0007"' in migration
    for table in (
        "steven_templates",
        "steven_tender_jobs",
        "steven_tender_suppliers",
        "steven_tender_versions",
        "steven_tender_candidate_links",
    ):
        assert table in migration
    assert "uq_steven_tender_supplier_name" in migration
    assert "uq_steven_tender_versions_job_version" in migration
    assert "steven:tenders:approve" in migration


def test_0012_migration_adds_batch_supplier_export_metadata_without_new_version_table():
    migration = (
        Path(__file__).resolve().parents[1]
        / "alembic"
        / "versions"
        / "20260717_0012_steven_s1_multi_supplier_exports.py"
    ).read_text(encoding="utf-8")
    assert 'revision = "20260717_0012"' in migration
    assert 'down_revision = "20260717_0011"' in migration
    assert "export_batch_id" in migration
    assert "supplier_name_snapshot" in migration
    assert "fk_steven_tender_versions_supplier" in migration
    assert "ck_steven_tender_versions_batch_supplier_metadata" in migration
    assert "supplier_name_snapshot IS NOT NULL" in migration
    assert "btrim(export_batch_id) <> ''" in migration
    assert "op.create_table" not in migration


def test_s1_demo_seed_uses_current_user_and_role_status_columns():
    seed = (
        Path(__file__).resolve().parents[3]
        / "scripts"
        / "seed_s1_demo.py"
    ).read_text(encoding="utf-8")
    assert "u.is_active" not in seed
    assert "JOIN roles r ON r.id=ur.role_id" in seed
    assert "r.code='operator'" in seed
    assert "r.status='active'" in seed
    assert "u.status='active'" in seed


@pytest.mark.parametrize(
    ("query", "expected_code"),
    [
        ("采购服务邀请", "DEMO-SERVICE-INVITATION"),
        ("quotation request", "DEMO-QUOTATION-REQUEST"),
        ("服务建议计划征集", "DEMO-SERVICE-PROPOSAL-REQUEST"),
    ],
)
def test_template_recommendation_ranks_matching_demo_template_first(tmp_path, query, expected_code):
    application, _, _ = make_application(tmp_path)

    recommendations = application.recommend_templates(query)

    assert recommendations
    assert recommendations[0].code == expected_code
    assert len(recommendations) <= 3
    assert application.recommend_templates("完全不相关的演示主题") == []


def test_template_recommendation_api_returns_keywords():
    client = TestClient(
        create_app(
            Settings(app_env="test", auth_mode="mock", demo_seed_enabled=True, mock_identity="steven"),
            InMemoryAuthRepository(),
        )
    )

    response = client.get("/api/v1/steven/tenders/templates/recommend", params={"q": "询价 request"})

    assert response.status_code == 200
    records = response.json()["data"]
    assert records[0]["code"] == "DEMO-QUOTATION-REQUEST"
    assert "询价" in records[0]["keywords"]


def test_approved_print_summary_is_supplier_isolated_and_escaped(tmp_path):
    application, _, _ = make_application(tmp_path)
    supplier_names = ["演示供应商甲", "演示供应商乙", "演示供应商<script>alert(1)</script>"]
    tender = create_tender(application, supplier_names=supplier_names)
    approve_tender(application, tender.id)

    summary = application.print_summary(tender.id)

    sections = re.findall(r'<section class="supplier-page">(.*?)</section>', summary, flags=re.DOTALL)
    assert len(sections) == len(supplier_names)
    for supplier_name, section in zip(supplier_names, sections, strict=True):
        visible_text = unescape(re.sub(r"<[^>]+>", "", section))
        assert supplier_name in visible_text
        assert all(other_name not in visible_text for other_name in supplier_names if other_name != supplier_name)
    assert "<script>alert(1)</script>" not in summary
    assert "&lt;script&gt;alert(1)&lt;/script&gt;" in summary
    assert "@media print" in summary
    assert "page-break-after:always" in summary


def test_print_summary_rejects_unapproved_tender(tmp_path):
    application, _, _ = make_application(tmp_path)
    tender = create_tender(application)

    with pytest.raises(ApiError) as error:
        application.print_summary(tender.id)

    assert error.value.code == "formal_print_forbidden"


def test_0015_migration_adds_jsonb_keywords_and_updates_existing_template_codes():
    migration = (
        Path(__file__).resolve().parents[1]
        / "alembic"
        / "versions"
        / "20260718_0015_steven_template_keywords.py"
    ).read_text(encoding="utf-8")
    assert 'revision = "20260718_0015"' in migration
    assert 'down_revision = "20260717_0014"' in migration
    assert '"keywords"' in migration
    assert "postgresql.JSONB" in migration
    assert "WHERE code = :template_code" in migration
    assert "WHERE template_code" not in migration
    for template_code in (
        "DEMO-SERVICE-INVITATION",
        "DEMO-QUOTATION-REQUEST",
        "DEMO-SERVICE-PROPOSAL-REQUEST",
    ):
        assert template_code in migration


def test_unapproved_tender_can_be_deleted_and_audited(tmp_path):
    application, _, audit = make_application(tmp_path)
    tender = create_tender(application, document_number="DEMO-S1-DELETE-001")

    result = application.delete_tender(tender.id, "steven.test")

    assert result == {"id": tender.id, "deleted": True}
    with pytest.raises(ApiError) as error:
        application.get_tender(tender.id)
    assert error.value.status_code == 404
    assert any(event["action"] == "tender.delete" for event in audit.list_for_object(tender.id))


def test_approved_tender_cannot_be_deleted(tmp_path):
    application, _, _ = make_application(tmp_path)
    tender = create_tender(application, document_number="DEMO-S1-DELETE-APPROVED")
    approve_tender(application, tender.id)

    with pytest.raises(ApiError) as error:
        application.delete_tender(tender.id, "steven.test")

    assert error.value.status_code == 409
    assert error.value.code == "approved_record_delete_forbidden"
