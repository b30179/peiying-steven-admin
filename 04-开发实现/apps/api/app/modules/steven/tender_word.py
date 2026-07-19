from __future__ import annotations

from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile, ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


FORBIDDEN_PACKAGE_MARKERS = (
    "vbaproject.bin",
    "word/embeddings/",
    "word/activex/",
    "oleobject",
)


class TenderWordRenderer:
    def render_to(self, target: Path, *, job, suppliers: list, rendered_body: str, formal: bool, version_number: int | None = None) -> None:
        document = self._build_document(
            job=job,
            suppliers=suppliers,
            rendered_body=rendered_body,
            formal=formal,
            version_number=version_number,
        )
        document.save(target)

    def render_bytes(self, *, job, suppliers: list, rendered_body: str, formal: bool = False) -> bytes:
        buffer = BytesIO()
        document = self._build_document(
            job=job,
            suppliers=suppliers,
            rendered_body=rendered_body,
            formal=formal,
            version_number=None,
        )
        document.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def verify(path: Path) -> None:
        Document(path)
        try:
            with ZipFile(path) as package:
                names = [name.replace("\\", "/").lower() for name in package.namelist()]
                if any(marker in name for name in names for marker in FORBIDDEN_PACKAGE_MARKERS):
                    raise ValueError("forbidden_docx_active_content")
                for name in package.namelist():
                    normalized = name.replace("\\", "/").lower()
                    if not normalized.endswith(".rels"):
                        continue
                    relationships = package.read(name).decode("utf-8", errors="replace").lower()
                    if 'targetmode="external"' in relationships or "targetmode='external'" in relationships:
                        raise ValueError("forbidden_docx_external_relationship")
        except BadZipFile as error:
            raise ValueError("invalid_docx_package") from error

    def _build_document(self, *, job, suppliers: list, rendered_body: str, formal: bool, version_number: int | None):
        document = Document()
        section = document.sections[0]
        section.start_type = WD_SECTION.NEW_PAGE
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

        normal = document.styles["Normal"]
        normal.font.name = "Microsoft YaHei"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        normal.font.size = Pt(10.5)

        banner = document.add_paragraph()
        banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
        banner.paragraph_format.space_after = Pt(10)
        banner_run = banner.add_run("正式批准版本" if formal else "草稿／未批准／仅供人工复核")
        self._set_run_font(banner_run, "Microsoft YaHei", 12, bold=True)
        self._shade_paragraph(banner, "E9F6EE" if formal else "FFF3D7")

        heading = document.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.paragraph_format.space_before = Pt(8)
        heading.paragraph_format.space_after = Pt(14)
        self._set_run_font(heading.add_run(job.title), "Microsoft YaHei", 20, bold=True)

        metadata = document.add_table(rows=0, cols=2)
        metadata.style = "Table Grid"
        metadata.autofit = True
        for label, value in (
            ("文书编号", job.document_number),
            ("事项名称", job.subject),
            ("生成日期", job.generated_date.isoformat()),
            ("截止日期", job.deadline_date.isoformat()),
            ("预算", f"{job.currency} {job.budget_min:.2f} - {job.budget_max:.2f}"),
            ("地点", job.location),
            ("版本", f"v{version_number}" if formal and version_number else "Draft"),
            ("资料标识", "当前使用脱敏演示资料"),
        ):
            cells = metadata.add_row().cells
            self._set_cell_text(cells[0], label, bold=True)
            self._set_cell_text(cells[1], str(value))

        self._add_heading(document, "文书正文")
        for line in rendered_body.splitlines():
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.line_spacing = 1.35
            paragraph.paragraph_format.space_after = Pt(4)
            self._set_run_font(paragraph.add_run(line or " "), "Microsoft YaHei", 10.5)

        self._add_heading(document, "邀请供应商")
        supplier_table = document.add_table(rows=1, cols=2)
        supplier_table.style = "Table Grid"
        self._set_cell_text(supplier_table.rows[0].cells[0], "序号", bold=True, centered=True)
        self._set_cell_text(supplier_table.rows[0].cells[1], "供应商（脱敏演示）", bold=True, centered=True)
        for index, supplier in enumerate(suppliers, start=1):
            cells = supplier_table.add_row().cells
            self._set_cell_text(cells[0], str(index), centered=True)
            self._set_cell_text(cells[1], supplier.supplier_name)

        notice = document.add_paragraph()
        notice.paragraph_format.space_before = Pt(14)
        notice.paragraph_format.line_spacing = 1.25
        self._set_run_font(
            notice.add_run(
                "本文件为 Steven AI 行政助手本机独立脱敏 Demo 产物。"
                "AI/OCR 未联网；正式状态仅由人工提交与审批产生。"
            ),
            "Microsoft YaHei",
            9,
        )
        return document

    @staticmethod
    def _set_run_font(run, name: str, size: float, *, bold: bool = False) -> None:
        run.font.name = name
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
        run.font.size = Pt(size)
        run.font.bold = bold

    def _set_cell_text(self, cell, value: str, *, bold: bool = False, centered: bool = False) -> None:
        cell.text = ""
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
        self._set_run_font(paragraph.add_run(value), "Microsoft YaHei", 10, bold=bold)

    def _add_heading(self, document, text: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(14)
        paragraph.paragraph_format.space_after = Pt(7)
        self._set_run_font(paragraph.add_run(text), "Microsoft YaHei", 14, bold=True)

    @staticmethod
    def _shade_paragraph(paragraph, fill: str) -> None:
        properties = paragraph._element.get_or_add_pPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        properties.append(shading)
