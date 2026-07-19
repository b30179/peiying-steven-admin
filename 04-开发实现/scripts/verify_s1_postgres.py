from __future__ import annotations

from contextlib import contextmanager
from datetime import date
from io import BytesIO
import json
import os
from pathlib import Path
import sys
from uuid import uuid4

from docx import Document
from pydantic import ValidationError
from sqlalchemy import create_engine, inspect, text
from sqlalchemy.engine import make_url


PROJECT_ROOT = Path(__file__).resolve().parents[1]
API_ROOT = PROJECT_ROOT / "apps" / "api"
EXPECTED_DATABASE = "puiying_steven_demo"
EXPECTED_REVISION = "20260717_0012"
EXPECTED_TABLES = {
    "steven_templates",
    "steven_tender_jobs",
    "steven_tender_suppliers",
    "steven_tender_versions",
    "steven_tender_candidate_links",
}
EXPECTED_CONSTRAINTS = {
    "ck_steven_tender_jobs_deadline",
    "ck_steven_tender_jobs_budget",
    "ck_steven_tender_jobs_status",
    "uq_steven_tender_jobs_document_number",
    "uq_steven_tender_supplier_name",
    "uq_steven_tender_versions_job_version",
    "uq_steven_tender_versions_storage_key",
    "fk_steven_tender_versions_supplier",
    "ck_steven_tender_versions_batch_supplier_metadata",
}
EXPECTED_INDEXES = {
    "ix_steven_tender_versions_export_batch",
    "ix_steven_tender_versions_supplier",
}

sys.path.insert(0, str(API_ROOT))

from app.core.api_response import ApiError  # noqa: E402
from app.core.audit_context import reset_request_id, set_request_id  # noqa: E402
from app.core.config import Settings  # noqa: E402
from app.main import LazyPostgresTenderApplication, create_app  # noqa: E402
from app.modules.steven.tender_repository import DEMO_TEMPLATE, DEMO_TEMPLATES  # noqa: E402
from app.modules.steven.tender_schemas import TenderCreateRequest  # noqa: E402


def emit(step: str, status: str, **details: object) -> None:
    print(json.dumps({"step": step, "status": status, **details}, ensure_ascii=False, default=str))


@contextmanager
def request_id(value: str):
    token = set_request_id(value)
    try:
        yield
    finally:
        reset_request_id(token)


def payload(document_number: str, *, template_id: str = DEMO_TEMPLATE.id, **overrides) -> TenderCreateRequest:
    values = {
        "template_id": template_id,
        "title": "脱敏采购服务邀请文书",
        "document_number": document_number,
        "subject": "脱敏设施保养服务",
        "generated_date": "2026-07-17",
        "deadline_date": "2026-07-20",
        "budget_min": 1000,
        "budget_max": 3000,
        "currency": "HKD",
        "location": "演示地点 A（脱敏）",
        "controlled_clauses": "仅接受脱敏演示资料；最终内容必须由人工复核。",
        "supplier_names": ["演示供应商甲", "演示供应商乙", "演示供应商丙"],
        "is_demo": True,
    }
    values.update(overrides)
    return TenderCreateRequest(**values)


def document_text(document: Document) -> str:
    paragraphs = [item.text for item in document.paragraphs]
    table_cells = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    return "\n".join([*paragraphs, *table_cells])


def actor_for_role(engine, role_code: str) -> str:
    with engine.connect() as connection:
        actor = connection.execute(
            text(
                """
                SELECT u.id
                  FROM users u
                  JOIN user_roles ur ON ur.user_id=u.id
                  JOIN roles r ON r.id=ur.role_id
                 WHERE u.status='active'
                   AND r.status='active'
                   AND r.code=:role_code
                 ORDER BY u.created_at,u.id
                 LIMIT 1
                """
            ),
            {"role_code": role_code},
        ).scalar_one_or_none()
    if actor is None:
        raise RuntimeError(f"active_{role_code}_actor_missing")
    return actor


def schema_smoke(engine) -> None:
    schema = inspect(engine)
    missing_tables = sorted(EXPECTED_TABLES - set(schema.get_table_names()))
    version_indexes = {item["name"] for item in schema.get_indexes("steven_tender_versions")}
    missing_indexes = sorted(EXPECTED_INDEXES - version_indexes)
    with engine.connect() as connection:
        revision = connection.execute(text("SELECT version_num FROM alembic_version")).scalar_one()
        constraints = set(
            connection.execute(
                text(
                    """
                    SELECT conname
                      FROM pg_constraint
                     WHERE conrelid IN (
                        'steven_tender_jobs'::regclass,
                        'steven_tender_suppliers'::regclass,
                        'steven_tender_versions'::regclass
                     )
                    """
                )
            ).scalars()
        )
        permissions = set(
            connection.execute(
                text("SELECT code FROM permissions WHERE code LIKE 'steven:tenders:%'")
            ).scalars()
        )
    missing_constraints = sorted(EXPECTED_CONSTRAINTS - constraints)
    if revision != EXPECTED_REVISION or missing_tables or missing_constraints or missing_indexes or len(permissions) != 6:
        raise RuntimeError(
            f"schema_contract_failed revision={revision} tables={missing_tables} "
            f"constraints={missing_constraints} indexes={missing_indexes} permissions={sorted(permissions)}"
        )
    emit(
        "schema",
        "passed",
        revision=revision,
        tables=sorted(EXPECTED_TABLES),
        constraints=sorted(EXPECTED_CONSTRAINTS),
        indexes=sorted(EXPECTED_INDEXES),
        permission_count=len(permissions),
    )


def insert_unresolved_template(engine, actor: str, suffix: str) -> str:
    template_id = str(uuid4())
    with engine.begin() as connection:
        connection.execute(
            text(
                """
                INSERT INTO steven_templates
                    (id,code,version,name,document_type,template_body,variables,status,is_demo,created_by,created_at,updated_at)
                VALUES
                    (:id,:code,1,:name,'procurement_service_invitation',:body,
                     CAST(:variables AS jsonb),'active',true,:actor,now(),now())
                """
            ),
            {
                "id": template_id,
                "code": f"DEMO-S1-UNRESOLVED-{suffix}",
                "name": "未替换变量阻断模板（脱敏验证）",
                "body": DEMO_TEMPLATE.template_body + "\n待人工补充：{{unresolved_demo_value}}",
                "variables": json.dumps([*DEMO_TEMPLATE.variables, "unresolved_demo_value"], ensure_ascii=False),
                "actor": actor,
            },
        )
    return template_id


def delete_unresolved_fixture(engine, template_id: str) -> None:
    with engine.begin() as connection:
        connection.execute(
            text(
                "DELETE FROM steven_tender_candidate_links "
                "WHERE tender_job_id IN ("
                "SELECT id FROM steven_tender_jobs WHERE template_id=:template_id AND is_demo=true"
                ")"
            ),
            {"template_id": template_id},
        )
        connection.execute(
            text(
                "DELETE FROM steven_tender_review_candidates "
                "WHERE tender_job_id IN ("
                "SELECT id FROM steven_tender_jobs WHERE template_id=:template_id AND is_demo=true"
                ")"
            ),
            {"template_id": template_id},
        )
        connection.execute(
            text(
                "DELETE FROM steven_tender_versions "
                "WHERE tender_job_id IN ("
                "SELECT id FROM steven_tender_jobs WHERE template_id=:template_id AND is_demo=true"
                ")"
            ),
            {"template_id": template_id},
        )
        connection.execute(
            text("DELETE FROM steven_tender_jobs WHERE template_id=:template_id AND is_demo=true"),
            {"template_id": template_id},
        )
        connection.execute(
            text(
                "DELETE FROM steven_templates "
                "WHERE id=:template_id AND is_demo=true AND code LIKE 'DEMO-S1-UNRESOLVED-%'"
            ),
            {"template_id": template_id},
        )


def assert_validation_rules(suffix: str) -> None:
    cases = (
        {"deadline_date": "2026-07-19"},
        {"budget_min": -1},
        {"budget_min": 3001, "budget_max": 3000},
    )
    for index, overrides in enumerate(cases, start=1):
        try:
            payload(f"DEMO-S1-RULE-{suffix}-{index}", **overrides)
        except ValidationError:
            continue
        raise RuntimeError(f"validation_rule_not_blocked_{index}")
    emit("input_rules", "passed", invalid_date=True, negative_budget=True, inverted_budget=True)


def application_smoke(settings: Settings, engine) -> dict[str, object]:
    application = create_app(settings)
    if not isinstance(application.state.tender_application, LazyPostgresTenderApplication):
        raise RuntimeError("tender_repository_not_postgresql")
    steven_actor = actor_for_role(engine, "operator")
    approver_actor = actor_for_role(engine, "approver")
    suffix = uuid4().hex[:10]
    tender_application = application.state.tender_application
    tender_application.ensure_demo_template(steven_actor)

    expected_template_codes = {item.code for item in DEMO_TEMPLATES}
    templates = [item for item in tender_application.list_templates() if item.code in expected_template_codes]
    if len(templates) != 3:
        raise RuntimeError("demo_template_count_below_three")
    if len({item.document_type for item in templates}) != 3:
        raise RuntimeError("demo_template_document_types_not_distinct")
    if len({tuple(item.variables) for item in templates}) != 3:
        raise RuntimeError("demo_template_variable_contracts_not_distinct")

    template_preview_ids: list[str] = []
    for index, template in enumerate(templates, start=1):
        template_tender = tender_application.create_tender(
            payload(
                f"DEMO-S1-TEMPLATE-{suffix}-{index}",
                template_id=template.id,
                title=f"{template.name}预览验证",
            ),
            steven_actor,
        )
        template_preview = tender_application.preview(template_tender.id, steven_actor)
        if not template_preview.valid or not template_preview.tender.rendered_body:
            raise RuntimeError(f"demo_template_preview_failed_{template.code}")
        if "{{" in template_preview.tender.rendered_body:
            raise RuntimeError(f"demo_template_unresolved_variable_{template.code}")
        template_preview_ids.append(template_tender.id)
    emit(
        "demo_templates",
        "passed",
        template_count=len(templates),
        codes=sorted(item.code for item in templates),
        document_types=sorted(item.document_type for item in templates),
    )

    assert_validation_rules(suffix)
    try:
        tender_application.create_tender(
            payload(
                f"DEMO-S1-DUP-{suffix}",
                supplier_names=["Demo Vendor", "demo vendor"],
            ),
            steven_actor,
        )
    except ApiError as error:
        if error.code != "duplicate_supplier_name":
            raise
    else:
        raise RuntimeError("duplicate_supplier_not_blocked")

    unresolved_template_id = insert_unresolved_template(engine, steven_actor, suffix)
    try:
        unresolved = tender_application.create_tender(
            payload(f"DEMO-S1-UNRESOLVED-{suffix}", template_id=unresolved_template_id),
            steven_actor,
        )
        unresolved_preview = tender_application.preview(unresolved.id, steven_actor)
        if unresolved_preview.valid or unresolved_preview.tender.status != "draft_error":
            raise RuntimeError("unresolved_variable_did_not_enter_draft_error")
        try:
            tender_application.submit(unresolved.id, steven_actor)
        except ApiError as error:
            if error.code != "draft_not_ready":
                raise
        else:
            raise RuntimeError("draft_error_submit_not_blocked")
        try:
            tender_application.export(unresolved.id, steven_actor)
        except ApiError as error:
            if error.code != "formal_export_forbidden":
                raise
        else:
            raise RuntimeError("draft_error_export_not_blocked")
    finally:
        delete_unresolved_fixture(engine, unresolved_template_id)

    returned = tender_application.create_tender(payload(f"DEMO-S1-RETURN-{suffix}"), steven_actor)
    tender_application.preview(returned.id, steven_actor)
    tender_application.submit(returned.id, steven_actor)
    try:
        tender_application.return_for_revision(returned.id, approver_actor, " ")
    except ApiError as error:
        if error.code != "return_opinion_required":
            raise
    else:
        raise RuntimeError("empty_return_opinion_not_blocked")
    tender_application.return_for_revision(returned.id, approver_actor, "请补充脱敏服务范围说明")
    try:
        tender_application.export(returned.id, steven_actor)
    except ApiError as error:
        if error.code != "formal_export_forbidden":
            raise
    else:
        raise RuntimeError("returned_export_not_blocked")

    request = f"s1-smoke-{suffix}"
    with request_id(request):
        tender = tender_application.create_tender(payload(f"DEMO-S1-APPROVED-{suffix}"), steven_actor)
        preview = tender_application.preview(tender.id, steven_actor)
        if not preview.valid or preview.tender.status != "draft":
            raise RuntimeError("valid_preview_failed")
        draft_filename, draft_content = tender_application.draft_bytes(tender.id, steven_actor)
        draft_document = Document(BytesIO(draft_content))
        if not draft_document.paragraphs or "草稿" not in "\n".join(item.text for item in draft_document.paragraphs):
            raise RuntimeError("draft_docx_not_reopenable")
        tender_application.submit(tender.id, steven_actor)
        try:
            tender_application.approve(tender.id, steven_actor, "不得成功")
        except ApiError as error:
            if error.status_code != 403 or error.code != "self_approval_forbidden":
                raise
        else:
            raise RuntimeError("self_approval_not_blocked")
        tender_application.approve(tender.id, approver_actor, "独立审批人已人工核对脱敏文书")
        exports = [tender_application.export(tender.id, steven_actor) for _ in range(3)]
        selected_suppliers = tender.suppliers[:2]
        batch = tender_application.batch_export(
            tender.id,
            [item.id for item in selected_suppliers],
            steven_actor,
        )

    if len(batch.versions) != 2 or any(item.status != "ready" for item in batch.versions):
        raise RuntimeError("batch_export_not_ready")
    if len({item.storage_key for item in batch.versions}) != 2:
        raise RuntimeError("batch_export_storage_key_reused")
    if {item.export_batch_id for item in batch.versions} != {batch.batch_id}:
        raise RuntimeError("batch_export_id_mismatch")
    if {item.supplier_id for item in batch.versions} != {item.id for item in selected_suppliers}:
        raise RuntimeError("batch_export_supplier_metadata_mismatch")

    selected_names = {item.id: item.supplier_name for item in selected_suppliers}
    for version in batch.versions:
        path = tender_application.version_file(tender.id, version.version_number)
        document = Document(path)
        content = document_text(document)
        expected_name = selected_names[version.supplier_id]
        other_names = {name for supplier_id, name in selected_names.items() if supplier_id != version.supplier_id}
        if expected_name not in content or any(name in content for name in other_names):
            raise RuntimeError("batch_export_supplier_data_leak")
        if tender.controlled_clauses not in content or "{{" in content:
            raise RuntimeError("batch_export_clause_or_variable_contract_failed")
        if version.supplier_name_snapshot != expected_name:
            raise RuntimeError("batch_export_supplier_snapshot_mismatch")

    versions = tender_application.list_versions(tender.id)
    if [item.version_number for item in versions] != [1, 2, 3, 4, 5]:
        raise RuntimeError("export_versions_not_continuous")
    if len({item.storage_key for item in versions}) != 5:
        raise RuntimeError("export_storage_key_reused")
    if any(item.status != "ready" for item in versions):
        raise RuntimeError("export_version_not_ready")
    file_root = Path(settings.file_storage_root).resolve()
    for version in versions:
        path = tender_application.version_file(tender.id, version.version_number)
        if file_root not in path.resolve().parents:
            raise RuntimeError("export_path_outside_file_root")
        document = Document(path)
        if not document.paragraphs or document.paragraphs[0].text != "正式批准版本":
            raise RuntimeError("formal_docx_not_reopenable")
        if path.stat().st_size != version.size_bytes or not version.sha256 or len(version.sha256) != 64:
            raise RuntimeError("export_metadata_mismatch")

    audits = tender_application.list_audit_events(tender.id)
    required_actions = {
        "tender.create",
        "tender.preview",
        "tender.submit",
        "tender.approve",
        "tender.export_reserved",
        "tender.export",
        "tender.batch_export_reserved",
        "tender.batch_export",
    }
    if not required_actions.issubset({item["action"] for item in audits}):
        raise RuntimeError("persistent_tender_audit_incomplete")
    if not any(item["request_id"] == request for item in audits):
        raise RuntimeError("tender_audit_request_id_missing")

    restarted = create_app(settings)
    recovered = restarted.state.tender_application.get_tender(tender.id)
    recovered_versions = restarted.state.tender_application.list_versions(tender.id)
    recovered_audits = restarted.state.tender_application.list_audit_events(tender.id)
    if recovered.status != "approved" or len(recovered_versions) != 5 or len(recovered_audits) < len(required_actions):
        raise RuntimeError("restart_recovery_failed")
    if not restarted.state.quote_application.list_quotes():
        raise RuntimeError("s2_baseline_missing_after_s1")
    with engine.connect() as connection:
        session_count = connection.execute(text("SELECT count(*) FROM auth_sessions")).scalar_one()
        s2_counts = connection.execute(
            text(
                """
                SELECT
                    (SELECT count(*) FROM steven_quote_jobs),
                    (SELECT count(*) FROM steven_quote_items),
                    (SELECT count(*) FROM steven_quote_suppliers),
                    (SELECT count(*) FROM steven_quote_offer_lines)
                """
            )
        ).one()
        file_rows = connection.execute(
            text(
                """
                SELECT count(*)
                  FROM files
                 WHERE module='steven'
                   AND document_type='tender_document'
                   AND purpose='approved_word_export'
                   AND request_id=:request_id
                """
            ),
            {"request_id": request},
        ).scalar_one()
    if session_count < 1 or file_rows != 5:
        raise RuntimeError("session_or_file_metadata_persistence_failed")

    return {
        "tender_id": tender.id,
        "draft_filename": draft_filename,
        "tender_status": recovered.status,
        "version_numbers": [item.version_number for item in recovered_versions],
        "template_count": len(templates),
        "template_document_types": sorted(item.document_type for item in templates),
        "template_preview_ids": template_preview_ids,
        "batch_id": batch.batch_id,
        "batch_supplier_count": len(batch.versions),
        "batch_supplier_isolation": True,
        "docx_reopenable": True,
        "request_id": request,
        "audit_count": len(recovered_audits),
        "session_count": session_count,
        "s2_counts": list(s2_counts),
        "unresolved_fixture_cleaned": True,
        "returned_tender_id": returned.id,
    }


def main() -> int:
    settings = Settings.from_env()
    settings.validate()
    database_name = make_url(settings.database_url).database if settings.database_url else None
    if settings.app_env != "development" or settings.auth_mode != "session":
        raise RuntimeError("s1_smoke_requires_development_session")
    if settings.demo_seed_enabled:
        raise RuntimeError("s1_smoke_forbids_runtime_demo_seed")
    if database_name != EXPECTED_DATABASE:
        raise RuntimeError("s1_smoke_wrong_database")
    if settings.ocr_enabled or settings.ai_structuring_enabled:
        raise RuntimeError("s1_smoke_requires_offline_ai_ocr")

    engine = create_engine(settings.database_url, pool_pre_ping=True, future=True)
    schema_smoke(engine)
    result = application_smoke(settings, engine)
    emit("s1_postgres_closure", "passed", **result)
    emit(
        "verification",
        "passed",
        scope="local_redacted_s1_demo",
        real_postgres=True,
        in_memory_fallback=False,
        external_services_called=False,
        production_ready=False,
        controlled_trial_ready=False,
        d2_authorized=False,
    )
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as error:
        emit("verification", "failed", error_type=type(error).__name__, message=str(error))
        raise SystemExit(1)
